"""
Document Export Functionality.

Export generated content to DOCX and PDF formats with professional formatting.
"""

import os
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from loguru import logger


def create_docx_document(
    content: str,
    title: str = "Patent Document",
    citations: List[Dict[str, Any]] = None,
    metadata: Dict[str, Any] = None,
) -> Document:
    """
    Create a professionally formatted DOCX document.
    
    Args:
        content: Main document content
        title: Document title
        citations: List of citations to include
        metadata: Additional metadata (authors, date, etc.)
    
    Returns:
        python-docx Document object
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata if provided
    if metadata:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if "inventors" in metadata:
            inventors = ", ".join(metadata["inventors"])
            run = meta_para.add_run(f"Inventors: {inventors}\n")
            run.font.size = Pt(10)
        
        if "date" in metadata:
            run = meta_para.add_run(f"Date: {metadata['date']}\n")
            run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
    
    # Add generated date
    gen_date_para = doc.add_paragraph()
    gen_date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = gen_date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # Add main content
    _add_formatted_content(doc, content)
    
    # Add citations section if provided
    if citations and len(citations) > 0:
        doc.add_page_break()
        doc.add_heading("References", level=1)
        
        for i, citation in enumerate(citations, 1):
            patent_id = citation.get("patent_id", "Unknown")
            relevance = citation.get("relevance", 0.0)
            snippet = citation.get("text_snippet", "")
            
            # Add citation
            cite_para = doc.add_paragraph()
            cite_para.add_run(f"[{i}] ").bold = True
            cite_para.add_run(f"{patent_id} ")
            cite_para.add_run(f"(Relevance: {relevance:.2f})").font.italic = True
            
            # Add snippet if available
            if snippet:
                snippet_para = doc.add_paragraph()
                snippet_para.style = 'Quote'
                snippet_para.add_run(snippet).font.size = Pt(9)
                doc.add_paragraph()  # Spacing
    
    logger.info(f"Created DOCX document with {len(doc.paragraphs)} paragraphs")
    
    return doc


def _add_formatted_content(doc: Document, content: str):
    """
    Add content to document with formatting preserved.
    
    Args:
        doc: Document object
        content: Content string
    """
    current_list_level = 0
    
    for line in content.split("\n"):
        line_stripped = line.strip()
        
        # Skip empty lines
        if not line_stripped:
            doc.add_paragraph()
            continue
        
        # Handle headings
        if line_stripped.startswith("###"):
            heading_text = line_stripped.strip("#").strip()
            doc.add_heading(heading_text, level=3)
        elif line_stripped.startswith("##"):
            heading_text = line_stripped.strip("#").strip()
            doc.add_heading(heading_text, level=2)
        elif line_stripped.startswith("#"):
            heading_text = line_stripped.strip("#").strip()
            doc.add_heading(heading_text, level=1)
        
        # Handle lists
        elif line_stripped.startswith("-") or line_stripped.startswith("*"):
            list_text = line_stripped[1:].strip()
            para = doc.add_paragraph(list_text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.5)
        
        elif line_stripped[0].isdigit() and line_stripped[1:3] in [". ", ") "]:
            list_text = line_stripped[2:].strip()
            para = doc.add_paragraph(list_text, style='List Number')
            para.paragraph_format.left_indent = Inches(0.5)
        
        # Regular paragraph
        else:
            para = doc.add_paragraph(line_stripped)
            para.paragraph_format.line_spacing = 1.15


def export_to_docx(
    content: str,
    output_path: str,
    title: str = "Patent Document",
    citations: List[Dict[str, Any]] = None,
    metadata: Dict[str, Any] = None,
) -> str:
    """
    Export content to a DOCX file.
    
    Args:
        content: Main document content
        output_path: Path to save DOCX file
        title: Document title
        citations: List of citations
        metadata: Additional metadata
    
    Returns:
        Absolute path to saved file
    
    Raises:
        RuntimeError: If export fails
    """
    try:
        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Create document
        doc = create_docx_document(content, title, citations, metadata)
        
        # Save document
        doc.save(output_path)
        
        abs_path = os.path.abspath(output_path)
        file_size = os.path.getsize(abs_path) / 1024  # KB
        
        logger.info(f"Exported DOCX to {abs_path} ({file_size:.1f} KB)")
        
        return abs_path
        
    except Exception as e:
        logger.error(f"Failed to export DOCX: {e}")
        raise RuntimeError(f"Failed to export document: {str(e)}")


def export_memo_to_docx(
    memo_text: str,
    citations: List[Dict[str, Any]],
    output_path: str = "invention_memo.docx",
    invention_title: str = None,
) -> str:
    """
    Export invention memo to DOCX.
    
    Args:
        memo_text: Memo content
        citations: Patent citations
        output_path: Where to save file
        invention_title: Optional invention title
    
    Returns:
        Path to saved file
    """
    title = invention_title or "Invention Disclosure Memo"
    
    metadata = {
        "date": datetime.now().strftime("%Y-%m-%d"),
        "inventors": ["To be determined"],
    }
    
    return export_to_docx(
        content=memo_text,
        output_path=output_path,
        title=title,
        citations=citations,
        metadata=metadata,
    )


def export_draft_to_docx(
    draft_text: str,
    citations: List[Dict[str, Any]],
    output_path: str = "patent_draft.docx",
    patent_title: str = None,
    inventors: List[str] = None,
) -> str:
    """
    Export patent draft to DOCX.
    
    Args:
        draft_text: Draft content
        citations: Patent citations
        output_path: Where to save file
        patent_title: Patent title
        inventors: List of inventor names
    
    Returns:
        Path to saved file
    """
    title = patent_title or "Patent Application Draft"
    
    metadata = {
        "date": datetime.now().strftime("%Y-%m-%d"),
    }
    
    if inventors:
        metadata["inventors"] = inventors
    
    return export_to_docx(
        content=draft_text,
        output_path=output_path,
        title=title,
        citations=citations,
        metadata=metadata,
    )


# ============================================================================
# PDF Export Functions
# ============================================================================

def create_pdf_document(
    content: str,
    output_path: str,
    title: str = "Patent Document",
    citations: List[Dict[str, Any]] = None,
    metadata: Dict[str, Any] = None,
) -> str:
    """
    Create a professionally formatted PDF document.
    
    Args:
        content: Main document content
        output_path: Path to save PDF file
        title: Document title
        citations: List of citations to include
        metadata: Additional metadata (authors, date, etc.)
    
    Returns:
        Absolute path to saved file
    
    Raises:
        RuntimeError: If PDF creation fails
    """
    try:
        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='black',
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Heading styles
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='black',
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='black',
            spaceAfter=10,
            spaceBefore=10,
            fontName='Helvetica-Bold'
        )
        
        # Body style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            textColor='black',
            spaceAfter=12,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        # Add title
        elements.append(Paragraph(title, title_style))
        elements.append(Spacer(1, 0.2 * inch))
        
        # Add metadata
        if metadata:
            meta_style = ParagraphStyle(
                'Metadata',
                parent=styles['Normal'],
                fontSize=10,
                textColor='grey',
                alignment=TA_CENTER
            )
            
            if "inventors" in metadata:
                inventors = ", ".join(metadata["inventors"])
                elements.append(Paragraph(f"Inventors: {inventors}", meta_style))
            
            if "date" in metadata:
                elements.append(Paragraph(f"Date: {metadata['date']}", meta_style))
            
            elements.append(Spacer(1, 0.3 * inch))
        
        # Add generation timestamp
        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=9,
            textColor='grey',
            alignment=TA_CENTER
        )
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M')
        elements.append(Paragraph(f"Generated: {timestamp}", timestamp_style))
        elements.append(Spacer(1, 0.5 * inch))
        
        # Add content
        for line in content.split('\n'):
            line = line.strip()
            
            if not line:
                elements.append(Spacer(1, 0.1 * inch))
                continue
            
            # Escape special characters for reportlab
            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Detect headings
            if line.startswith('###'):
                text = line.strip('#').strip()
                elements.append(Paragraph(text, heading2_style))
            elif line.startswith('##'):
                text = line.strip('#').strip()
                elements.append(Paragraph(text, heading1_style))
            elif line.startswith('#'):
                text = line.strip('#').strip()
                elements.append(Paragraph(text, title_style))
            # Detect lists
            elif line.startswith('-') or line.startswith('*'):
                text = '• ' + line[1:].strip()
                elements.append(Paragraph(text, body_style))
            # Regular paragraph
            else:
                elements.append(Paragraph(line, body_style))
        
        # Add citations if provided
        if citations and len(citations) > 0:
            elements.append(PageBreak())
            elements.append(Paragraph("References", heading1_style))
            elements.append(Spacer(1, 0.2 * inch))
            
            citation_style = ParagraphStyle(
                'Citation',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=10,
                leftIndent=20
            )
            
            for i, citation in enumerate(citations, 1):
                patent_id = citation.get("patent_id", "Unknown")
                relevance = citation.get("relevance", 0.0)
                snippet = citation.get("text_snippet", "")
                
                cite_text = f"[{i}] {patent_id} (Relevance: {relevance:.2f})"
                elements.append(Paragraph(cite_text, citation_style))
                
                if snippet:
                    snippet_style = ParagraphStyle(
                        'Snippet',
                        parent=styles['Normal'],
                        fontSize=9,
                        textColor='grey',
                        leftIndent=40,
                        spaceAfter=10
                    )
                    elements.append(Paragraph(snippet, snippet_style))
        
        # Build PDF
        doc.build(elements)
        
        abs_path = os.path.abspath(output_path)
        file_size = os.path.getsize(abs_path) / 1024  # KB
        
        logger.info(f"Exported PDF to {abs_path} ({file_size:.1f} KB)")
        
        return abs_path
        
    except Exception as e:
        logger.error(f"Failed to export PDF: {e}")
        raise RuntimeError(f"Failed to export PDF: {str(e)}")


def export_memo_to_pdf(
    memo_text: str,
    citations: List[Dict[str, Any]],
    output_path: str = "invention_memo.pdf",
    invention_title: str = None,
) -> str:
    """
    Export invention memo to PDF.
    
    Args:
        memo_text: Memo content
        citations: Patent citations
        output_path: Where to save file
        invention_title: Optional invention title
    
    Returns:
        Path to saved file
    """
    title = invention_title or "Invention Disclosure Memo"
    
    metadata = {
        "date": datetime.now().strftime("%Y-%m-%d"),
        "inventors": ["To be determined"],
    }
    
    return create_pdf_document(
        content=memo_text,
        output_path=output_path,
        title=title,
        citations=citations,
        metadata=metadata,
    )


def export_draft_to_pdf(
    draft_text: str,
    citations: List[Dict[str, Any]],
    output_path: str = "patent_draft.pdf",
    patent_title: str = None,
    inventors: List[str] = None,
) -> str:
    """
    Export patent draft to PDF.
    
    Args:
        draft_text: Draft content
        citations: Patent citations
        output_path: Where to save file
        patent_title: Patent title
        inventors: List of inventor names
    
    Returns:
        Path to saved file
    """
    title = patent_title or "Patent Application Draft"
    
    metadata = {
        "date": datetime.now().strftime("%Y-%m-%d"),
    }
    
    if inventors:
        metadata["inventors"] = inventors
    
    return create_pdf_document(
        content=draft_text,
        output_path=output_path,
        title=title,
        citations=citations,
        metadata=metadata,
    )

